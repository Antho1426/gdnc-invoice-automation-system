# Script name:         generate_sport_invoices.py
# Python interpreter:  Miniconda virtual environment "automation-env"
# Description:         Code for automatically generating and sending GDNC sport invoices to people having registered on our website (https://concise2025.ch/inscriptions-sports/) and which data can be read from Excel file
# Invocation example:  python generate_sport_invoices.py
# Author:              Anthony Guinchard
# Version:             0.1
# Creation date:       2025-06-20
# Modification date:   2025-06-20
# Resources:
#                      - YouTube video "NeuralNine - Invoice Automation System in Python - Full Project" (https://youtu.be/JuBEC1RW8nA?si=-a1BploFfwDJsV0a).
#                      - Discussion with ChatGPT (https://chatgpt.com/c/6792c47d-01b4-8003-abc4-23d175330cdc)
# Working:             ✅

import json
import math as m
import os
import platform
import re
import subprocess
import sys
import threading
import time
from halo import Halo
import tkinter as tk
from datetime import datetime, timedelta
from enum import Enum
from functools import wraps
from pathlib import Path
from openpyxl import load_workbook
from time import perf_counter
from typing import Dict, Any
from tkinter import (BOTH, LEFT, RIGHT, VERTICAL, Canvas, Frame, Y, filedialog,
                     messagebox, ttk)
import pyautogui
from pynput.keyboard import Controller


import click
import docx
import pandas as pd
import pyperclip
import requests
import threading
from definition import (BIN_PATH, CURRENT_TIME, INVOICE_MODELS_FOLDER_NAME,
                        LIB_PATH, LOG_PATH, OUT_PATH, PROJECT_PATH,
                        SOFFICE_BINARY_PATH, SPONSOR_DATABASE_DEBUG_NAME, SPONSOR_DATABASE_NAME, SHEET_NAME, SPONSOR_DATABASE_DEBUG_NAME, SPONSOR_DATABASE_PATH,
                        SPORTS_CATALOG_PATH, SRC_PATH, NUM_INVOICE_PATH, DEBUG_MODE, REGISTRATION_EXCEL_FILE_NAME, SPORTS_SHEET_NAME_LIST, SPORTS_LIST)
from docx2pdf import convert
from pandas import DataFrame
from PIL import Image, ImageSequence, ImageTk
from termcolor import colored
from tktooltip import ToolTip

from selenium import webdriver
from selenium.webdriver.chrome.service import Service as ChromeService
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.webdriver import WebDriver
from selenium.common.exceptions import NoSuchElementException

from utils import (DualLogger, get_deadline_formatted_date, get_invoice_number,
                   get_today_formatted_date, replace_text)

SCRIPT_NAME = Path(__file__).name


def launch_client_chrome_instance():
    """Launch client Chrome instance in separate terminal thread via iTerm.
    """
    # Open iTerm app (will bring to front if already open)
    subprocess.Popen(["open", "-a", "iTerm"])

    # Wait for iTerm to open and be ready
    time.sleep(2)  # adjust if needed

    # Type command in iTerm
    cmd = "/Applications/Google\ Chrome.app/Contents/MacOS/Google\ Chrome --remote-debugging-port=8989 --user-data-dir=/Users/anthony/Library/Application Support/Google/Chrome"
    keyboard = Controller()
    keyboard.type(cmd)

    # Press Enter to run the command
    pyautogui.press("enter")


def setup_selenium():
    """Function to set up Selenium in order to interact with Infomaniak website.
    :return driver: The Chrome webdriver used to interact with the webpage.

    Options to connect to Infomaniak mailbox:

        Option 1: Use Infomaniak API → Doesn't work!
        ------------------------------------------------------------------------
        import smtplib
        from email.message import EmailMessage

        msg = EmailMessage()
        msg["From"] = "finances@concise2025.ch"
        msg["To"] = "antho.guinchard@gmail.com"
        msg["Subject"] = "Test GDNCInvoiceAutomationSystem"
        msg.set_content("Hello via Infomaniak SMTP!")

        with smtplib.SMTP_SSL("mail.infomaniak.com", 465) as smtp:
            #smtp.login("antho.guinchard@gmail.com", "1nf0_m4niAq_p0l1c317.7")
            smtp.login("finances@concise2025.ch", "JGe8kKVjauKwhNtg")
            smtp.send_message(msg)

        Option 2: Use Selenium → Works! See required steps and implementation in function below!
        ------------------------------------------------------------------------
        1) Enable CDP (Chrome DevTools Protocol):
            Prior to running the Python script scraping a website with Selenium, launch Chrome with Remote Debugging Port Enabled: Open your terminal and run the following command, replacing /path/to/chrome with the actual path to your Chrome executable and /path/to/any/directory/where/you/want/to/set/your/chrome/profile with the actual path where you want to start Chrome in a separate client Chrome instance:

            /Applications/Google\ Chrome.app/Contents/MacOS/Google\ Chrome --remote-debugging-port=8989 --user-data-dir=/path/to/any/directory/where/you/want/to/set/your/chrome/profile

            Adapted command for my use-case (directly pointing towards the location of my default Chrome profile in "Application Support"):

            /Applications/Google\ Chrome.app/Contents/MacOS/Google\ Chrome --remote-debugging-port=8989 --user-data-dir=/Users/anthony/Library/Application Support/Google/Chrome

            This command starts Chrome in a separate client Chrome instance, using a distinct user profile, and opens a debugging interface on port 8989. You can choose a different port if necessary. This is the Chrome instance where your selenium tests will be running. 

        2) Setting up Selenium with CDP:
            To run Selenium tests on an already opened Chrome browser, you'll need to set up a Selenium WebDriver instance with the CDP enabled. For secured website, like LinkedIn, make sure to identify yourself beforehand (i.e., by MANUALLY accessing the website from client Chrome instance with username and password!).
    """
    # Launch client Chrome instance in terminal
    launch_client_chrome_instance()

    # Specify the debugging address for the already opened Chrome browser
    debugger_address = "localhost:8989"

    # Set up ChromeOptions and connect to the existing browser
    c_options = webdriver.ChromeOptions()
    c_options.add_experimental_option("debuggerAddress", debugger_address)

    # Initialize the WebDriver with the existing Chrome instance
    service = ChromeService(ChromeDriverManager().install())
    driver = webdriver.Chrome(service=service, options=c_options)
    driver.set_window_size(1000, 850)

    # Now, you can interact with the already opened Chrome browser
    mailbox_url = "https://mail.infomaniak.com/0"
    driver.get(mailbox_url)

    # Wait 10 seconds until URL may be the mailbox URL (otherwise, it means that we land onto the login page)
    WebDriverWait(driver, 10).until(lambda d: d.current_url == mailbox_url)

    if driver.current_url != mailbox_url:
        # ❗ In order to access secured website, before continuing and beginning scraping, MANUALLY enter your credentials (username + password). The program will stop here until you manually input your credentials.
        input("👉 Before continuing, if necessary, please manually enter your credentials to connect to Infomaniak secured mailbox website in separate client Chrome instance window. Then press Enter to continue...")

    return driver


def shutdown_selenium(driver: WebDriver):
    driver.quit()


def click_button(driver: WebDriver, xpath: str) -> None:
    """Click on a button in a webpage using Selenium package and provided button XPath.
    
    :param driver: The Chrome webdriver used to interact with the webpage.
    :param xpath: Button XPath under the form of string.
    :return: None.
    """
    WebDriverWait(driver, 10).until(EC.element_to_be_clickable((By.XPATH, xpath))).click()


def enter_input(driver: WebDriver, xpath: str, text_input: str) -> None:
    """Enter a string in a text field of a webpage using Selenium package and provided input XPath.
    
    :param driver: The Chrome webdriver used to interact with the webpage.
    :param xpath: Input XPath under the form of string.
    :param text_input: The text to input in the input field.
    :return: None.
    """
    # Wait until input field is present and interactable
    input_elem = WebDriverWait(driver, 10).until(
        EC.visibility_of_element_located((By.XPATH, xpath))
    )
    # Paste a value into the input field
    input_elem.clear()  # clear existing text
    input_elem.send_keys(text_input)


def enter_editor(driver: WebDriver, editor_name: str, text_input: str) -> None:
    """Enter a string in an editor field of a webpage using Selenium package and provided editor name.
    
    :param driver: The Chrome webdriver used to interact with the webpage.
    :param editor_name: Name of editor element under the form of string.
    :param text_input: The text to input in the input field.
    :return: None.
    """
    editor = driver.find_element(By.ID, editor_name)
    editor.click()
    editor.clear()  # clear existing text
    editor.send_keys(text_input)


def compose_email(registrer_dict: dict) -> str:
    """Compose email to send based on current  day of the week and current hour
    :param registrer_dict: A dictionary containing all invoice-related fields,
        including customer information and products purchased.
    :param inscriptions: A list containing the information regarding the different sports registrations of the registrer.
    :return email_message: A string containing the email text to send.
    """
        
    current_day = datetime.now().weekday()  # Monday = 0, Tuesday = 1, Wednesday = 2, Thursday = 3, Friday = 4, Saturday = 5, Sunday = 6

    inscriptions = registrer_dict["custom product"]
    inscription_list = [next(iter(inscriptions[i].values()))['description'] for i in range(len(inscriptions))]
    inscription_list_formatted = "\n\n".join(f"    • {desc}" for desc in inscription_list)
    email_template = f"Salut cher sportif!\n\n\nMerci beaucoup pour ton intérêt aux sports du Giron du Nord 2025 à Concise! Suite à ta récente inscription, nous t'envoyons maintenant la facture conformément à ton enregistrement aux différentes disciplines.\n\n\nPour rappel, voici un résumé des sports auxquels tu t'es inscrits:\n\n{inscription_list_formatted}\n\n\nNous restons à ton entière disposition pour toute précision. Merci encore beaucoup pour ton enregistrement!\n\n\nMeilleures salutations et [CLOSING],\n\n\nAnthony Guinchard\n\nCommission Finances\n\nGiron du Nord 2025 à Concise"
    
    class Closing(Enum):
        START_OF_WEEK = "encore un bon début de semaine"
        MID_WEEK = "une bonne suite de semaine"
        END_OF_WEEK = "déjà une bonne fin de semaine"
        WEEKEND = "un bon weekend"
        END_OF_WEEKEND = "déjà un bon début de semaine"

    match current_day:
        case 0 | 1:  # Monday or Thursday
            closing = Closing.START_OF_WEEK.value
        case 2 | 3: # Wednesday or Tuesday
            closing = Closing.MID_WEEK.value
        case 4: # Friday
            closing = Closing.END_OF_WEEK.value
        case 5: # Saturday
            closing = Closing.WEEKEND.value
        case 6: # Sunday
            closing = Closing.END_OF_WEEKEND.value
    
    email_message = email_template.replace("[CLOSING]", closing)
    
    print(f"\t\t\t▷ Generated email message:\n---\n{email_message}\n---")

    return email_message


def load_registrations_from_excel() -> DataFrame:
    """Load registration data from an Excel file and return it as a pandas DataFrame.

    This function performs the following checks:
    
    - Verifies the existence of the Excel file at the specified path.
    - Ensures the file is not empty.
    - Checks for the presence of all required columns.

    If any check fails, the function prints an error and terminates the program.

    :returns: A pandas DataFrame containing the registration data.
    :rtype: pandas.DataFrame
    :raises SystemExit: If the file does not exist, is empty, or is missing required columns.
    """
    excel_file_path = LIB_PATH / REGISTRATION_EXCEL_FILE_NAME
    if not excel_file_path.exists():
        print(colored("Error!", "red"), f"The file '{excel_file_path}' does not exist. Program will stop here.")
        sys.exit(1)
    
    # Read the different sheets from the Excel file
    print(f"Reading data from '{excel_file_path}'...")
    df_list = []
    for sheet_name in SPORTS_SHEET_NAME_LIST:
        print(f"\tReading sheet '{sheet_name}'...")
        try:
            df = pd.read_excel(excel_file_path, sheet_name=sheet_name)
            print(f"\t\tData read successfully from '{sheet_name}' sheet.")
            # Add a column for the sport name
            df["Sport"] = SPORTS_LIST[SPORTS_SHEET_NAME_LIST.index(sheet_name)]
            df_list.append(df)
        except ValueError as e:
            print(colored("\t\tError!", "red"), f"Failed to read sheet '{sheet_name}': \n\t\t\t{e}\n\t\t\tProgram will stop here.")
            sys.exit(1)
    
    # Concatenate all sheets into a single DataFrame
    df = pd.concat(df_list, ignore_index=True)
    # Check if the DataFrame is empty
    if df.empty:
        print("No registrations found in the Excel file. Program will stop here.")
        sys.exit(1)
    
    # Check if the required columns are present
    required_columns = ["Entry ID", "Date Created", "Nom complet", "E-mail", "Téléphone", "Adresse", "Nom d'équipe", "Nombre d'équipe(s)", "Total", "Sport"]
    for column in required_columns:
        if column not in df.columns:
            print(colored("\t\tError!", "red"), f"The required column '{column}' is missing in the Excel file. Program will stop here.")
            sys.exit(1)
    
    print(f"Data read successfully. Total number of individual teams registered: {len(df)}")

    return df


def sanitize_data(df: DataFrame) -> DataFrame:
    """Sanitize and transform registration data.

    This function processes a registration DataFrame to:
    - Deduplicate registrants based on unique email addresses
    - Extract and group registered sports and associated teams
    - Generate current registration and deadline dates
    - Parse and validate address fields (street, postcode, city)
    - Build and return a cleaned DataFrame with relevant fields

    :param df: Raw registration DataFrame.
    :return df_sanitized: A sanitized and structured DataFrame with one row per registrant,
        including sports data, formatted dates, and parsed address fields.
    :raises SystemExit: If address format errors are detected.
    """
    print("Sanitizing the data...")
    
    # Registered sports
    email_unique_list = df["E-mail"].unique().tolist()
    num_registrers = len(email_unique_list)
    # Gather sports and team names for each unique email
    registered_sports_list = []
    for email in email_unique_list:
        df_registrer = df[df["E-mail"] == email]
        sports_unique_list = df_registrer["Sport"].unique().tolist()
        global_sport_dict = {}
        for sport in sports_unique_list:
            sport_df = df_registrer[df_registrer["Sport"] == sport]
            team_name_list = sport_df["Nom d'équipe"].tolist()
            sport_dict = {sport: team_name_list}
            global_sport_dict.update(sport_dict)
        registered_sports_list.append(global_sport_dict)
        
    # Dates
    # From "Date Created" column
    # date_created_list = df["Date Created"].unique().tolist()
    # date_list = [date.split(" ")[0].replace("-",".") for date in date_created_list]
    # date_obj_list = [datetime.strptime(date, "%Y.%m.%d") for date in date_list]
    # date_list = [date_obj.strftime("%d.%m.%Y") for date_obj in date_obj_list]
    # From current invoicing date
    date_obj_list = [datetime.now() for _ in range(num_registrers)]  # use current date for all entries
    date_list = [get_today_formatted_date() for _ in range(num_registrers)]
    date_deadline_list = [get_deadline_formatted_date(date_obj) for date_obj in date_obj_list]

    # Address
    address_unique_list = (
        df[df["E-mail"].isin(email_unique_list)]
        .drop_duplicates(subset="E-mail", keep="first")["Adresse"]
        .tolist()
    )    
    # Street
    street_list = [address.split(",")[0].strip() for address in address_unique_list]
    # Postcode
    postcode_list = [address.split(",")[2].strip() for address in address_unique_list]
    # City
    city_list = [address.split(",")[1].strip() for address in address_unique_list]
    # Eventually exchange postcode and city where necessary
    for i in range(len(postcode_list)):
        postcode = postcode_list[i]
        city = city_list[i]
        if postcode.isalpha() and city.isdigit():
                postcode_list[i], city_list[i] = city_list[i], postcode_list[i]
        elif postcode.isalpha() and city.isalpha():
            print(colored("\tError!", "red"), f"Postcode '{postcode}' and city '{city}' are both alphabetic for registrer {email_unique_list[i]}. Please check the address format. Program will stop here.")
            sys.exit(1)
        elif postcode.isdigit() and city.isdigit():
            print(colored("\tError!", "red"), f"Postcode '{postcode}' and city '{city}' are both numeric for registrer {email_unique_list[i]}. Please check the address format. Program will stop here.")
            sys.exit(1)
        else:
            continue

    # Phone
    phone_unique_list = (
        df[df["E-mail"].isin(email_unique_list)]
        .drop_duplicates(subset="E-mail", keep="first")["Téléphone"]
        .tolist()
    )
    phone_list = [str(phone) for phone in phone_unique_list]
    phone_list_adjusted = []
    for phone in phone_list:
        if phone.startswith("0"):
            phone = "+41" + phone[1:]
        if phone.startswith("41"):
            phone = "+" + phone
        if phone == "nan":
            phone = ""
        phone_list_adjusted.append(phone)
 
    # Build sanitized DataFrame
    # Retrieve old columns to keep
    columns_to_keep = ["Entry ID", "Nom complet", "Téléphone", "E-mail"]
    df_sanitized = (
        df[df["E-mail"].isin(email_unique_list)]
        .drop_duplicates(subset="E-mail", keep="first")[columns_to_keep]
    )
    # Rename some of the old columns in English
    df_sanitized.rename(columns={
        "Nom complet": "Name",
        "Téléphone": "Phone",
        "E-mail": "Email",
    }, inplace=True)
    # Add new columns
    df_sanitized["Registered Sports"] = registered_sports_list
    df_sanitized["Date"] = date_list
    df_sanitized["Deadline"] = date_deadline_list
    df_sanitized["Street"] = street_list
    df_sanitized["Postcode"] = postcode_list
    df_sanitized["City"] = city_list
    df_sanitized["Phone"] = phone_list_adjusted
    # Reset index
    df_sanitized = df_sanitized.reset_index(drop=True)

    print(f"Data sanitized successfully. Found {num_registrers} registrers.")

    return df_sanitized


def generate_invoice(entry: Dict[str, Any]) -> Dict[str, Any]:
    """Generate a personalized invoice document (DOCX and PDF) for a sports
    registration entry, update tracking files, and return a summary dictionary
    of the invoice.

    This function:
      - Computes pricing based on sports and teams registered
      - Fills in a DOCX invoice template with participant and invoice details
      - Converts the invoice to PDF using LibreOffice
      - Updates a tracking Excel file with invoice metadata
      - Returns a dictionary used to update the invoice database

    :param entry: A dictionary containing participant and registration data.
    :type entry: dict[str, Any]
    :return registrer_dict: A dictionary summarizing invoice data to be used for
        database updates.
    :return invoice_path: A string containing the path to the generated PDF invoice.
    """
    # Get invoice number
    invoice_number = get_invoice_number()
    
    print(f"\t\tProcess launched for generating invoice {invoice_number}! 🚀")
    time_start = perf_counter() 

    # Generate DOCX document

    # Compute price
    num_total_products = len(entry["Registered Sports"])
    product_dict_list = []
    for key, value in entry["Registered Sports"].items():
        sport = key
        sport_price = next((v["price"] for v in SPORTS_CATALOG_DICT.values() if v["name"] == sport), None)
        num_teams = len(value)
        registration = "Inscriptions" if num_teams > 1 else "Inscription"
        team = "équipes" if num_teams > 1 else "équipe"
        price = num_teams*sport_price
        description = f"{registration} {sport} ({team}: {', '.join(value)})"
        product_dict = {sport: {"description": description, "num teams": num_teams, "price": price}}
        product_dict_list.append(product_dict)
    # Compute total price
    total_price = 0
    for product_dict in product_dict_list:
        price = next(iter(product_dict.values()))["price"]
        total_price += price

    # Update status label
    print("\t\t\t> Select invoice template...")
    doc = docx.Document(template_path)

    product_key = "[PRODUCT-DESCRIPTION-IDX]"
    quantity_key = "[QT-IDX]"
    price_key = "[P-IDX]"
    tot_key = "[TOT-IDX]"

    replacements = {
        "[COMPANY]": "",  # no company for sports
        "[TITLE]": "",  # no title for sports
        "[FIRST-NAME]": "",  # no first name for sports
        "[LAST-NAME]": entry["Name"],
        "[ADDRESS]": entry["Street"],
        "[POSTCODE]": entry["Postcode"],
        "[CITY]": entry["City"],
        "[INVOICE-NUMBER]": invoice_number,
        "[ISSUE-DATE]": entry["Date"],
        "[DEADLINE-DATE]": entry["Deadline"],
        "[TOTAL]": str("{:.2f}".format(float(total_price)))  # convert price to string since all mapped values have to have type string
    }

    for i, product_dict in enumerate(product_dict_list, start=1):
        product_replacements = {}
        
        sport = next(iter(product_dict.keys()))
        name = next(iter(product_dict.values()))["description"]
        quantity = next(iter(product_dict.values()))["num teams"]
        price = next(iter(product_dict.values()))["price"]
        idx = i
        
        product_key_idx = product_key.replace("IDX", str(idx))
        product_replacements[product_key_idx] = name
        quantity_key_idx = quantity_key.replace("IDX", str(idx))
        product_replacements[quantity_key_idx] = str(quantity)
        price_key_idx = price_key.replace("IDX", str(idx))
        sport_price = next((v["price"] for v in SPORTS_CATALOG_DICT.values() if v["name"] == sport), None)
        product_replacements[price_key_idx] = str(sport_price)  # convert price to string since all mapped values have to have type string
        tot_key_idx = tot_key.replace("IDX", str(idx))
        product_replacements[tot_key_idx] = str("{:.2f}".format(float(price)))

        replacements.update(product_replacements)

    # Update status label
    print("\t\t\t> Replace keys in template...")
 
    # Make replacements in the paragraphs of the DOCX document (i.e., info and invoice data)
    for paragraph in list(doc.paragraphs):
        for old_text, new_text in replacements.items():
            replace_text(paragraph=paragraph,
                            old_text=old_text, new_text=new_text)

    # Make replacements in the tables of the DOCX document (i.e., product data)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        replace_text(
                            paragraph=paragraph, old_text=old_text, new_text=new_text)
                    # Make sure product number "01", "02", etc. (under the column "NO" in the invoice template) are not bold
                    if re.fullmatch(r"0[1-5]", paragraph.text):
                        paragraph.style.font.bold = False
                        for run in paragraph.runs:
                            run.bold = False

    invoice_name = f"Facture N° {invoice_number}.docx"
    if DEBUG_MODE:
        invoice_name = invoice_name.replace(".docx", "_DEBUG.docx")
    output_docx_path = str(OUT_PATH / invoice_name)
    doc.save(output_docx_path)
    invoice_path = output_docx_path.replace(".docx",".pdf")
    
    # Convert DOCX to PDF
    
    print("\t\t\t> DOCX to PDF conversion...")

    if not SOFFICE_BINARY_PATH.exists():
        print(f"LibreOffice binary file '{SOFFICE_BINARY_PATH}' does not exist. Please download LibreOffice to your Mac from 'https://www.libreoffice.org/donate/dl/mac-x86_64/25.2.1/fr/LibreOffice_25.2.1_MacOS_x86-64.dmg' or, if using Linux operating system, install it using the command `sudo apt install libreoffice` (in this case, make sure to add line `export LD_LIBRARY_PATH=/usr/lib/libreoffice/program:$LD_LIBRARY_PATH` to your .bashrc and .zshrc files to avoid issues such as `/usr/lib/libreoffice/program/soffice.bin: error while loading shared libraries: libreglo.so: cannot open shared object file: No such file or directory`) or download the Debian file from 'https://www.libreoffice.org/download/download-libreoffice/?type=deb-x86_64&version=25.2.1&lang=en-US'. The DOCX invoice could be generated but not converted into PDF. Invoice generation will stop here.")
        sys.exit(1)

    spinner = Halo(text="", spinner='dots')
    spinner.start()

    try:
        # GUI solution for converting DOCX to PDF (using Microsoft Word) (not reliable every time; produces errors like: "'result': 'error', 'error': 'Error: Message not understood.'")
        #convert(input_path=output_docx_path, output_path=output_pdf_path)
        # Headless solution for converting DOCX to PDF (using LibreOffice with command `soffice --headless --convert-to pdf:writer_pdf_Export --outdir out/ input.docx`) (see "https://github.com/AlJohri/docx2pdf/issues/51#issuecomment-1335382983" and "https://stackoverflow.com/a/32595547") (download LibreOffice for macOS from this link: https://www.libreoffice.org/donate/dl/mac-x86_64/25.2.1/fr/LibreOffice_25.2.1_MacOS_x86-64.dmg)
        subprocess.run([SOFFICE_BINARY_PATH, "--headless", "--convert-to", "pdf:writer_pdf_Export", "--outdir", str(OUT_PATH), output_docx_path], check=True)
        spinner.succeed()
        print(f"\t\t\t\tDOCX to PDF conversion successful!")
    except Exception as e:
        spinner.fail()
        print(colored("\t\t\t\tError!", "red"), f"DOCX to PDF conversion failed:\n\t\t\t\t\t{e}.\n\t\t\t\t\tProgram will stop here.")

    time_end = perf_counter()
    elapsed_time = time_end - time_start
    print(f"\t\t\tInvoice created and saved successfully! ⏱️ Elapsed time: {elapsed_time:.2f} [s]")

    # Generate new line to fill in file "1_N° facture.xlsx"
    num_invoice_entry_list = [get_today_formatted_date().replace('.', '/'), invoice_number, f"{entry['Name']} (sports)", str(int(total_price)), "Mail"]
    print(f"\t\t▷ Generated new line for file '1_N° facture.xlsx':\n\t\t\t{num_invoice_entry_list}")
    workbook = load_workbook(NUM_INVOICE_PATH)
    # Select the sheet
    sheet = workbook["Facturation"]
    # Append the row at the very bottom of the table
    # Find last non-empty row based on a key column
    last_row = 1
    for row in range(2, sheet.max_row + 1):
        if sheet.cell(row=row, column=1).value not in (None, ""):
            last_row = row
    # Write your data in the next row
    for col_index, value in enumerate(num_invoice_entry_list, start=1):
        sheet.cell(row=last_row + 1, column=col_index, value=value)
    # Save changes
    workbook.save(NUM_INVOICE_PATH)

    registrer_dict = {
        "date": get_today_formatted_date(),
        "invoice number": get_invoice_number(),
        "company": "",
        "title": "",
        "first name": "",
        "last name": entry["Name"],
        "address": entry["Street"],
        "postcode": entry["Postcode"],
        "city": entry["City"],
        "phone": entry["Phone"],
        "email": entry["Email"],
        "default product": "",
        "custom product": product_dict_list,
        "total price": total_price,
        "comment": "",
    }

    return registrer_dict, invoice_path


def update_invoice_database(registrer_dict: Dict[str, Any]) -> None:
    """Backup registerer data and update invoice database with a new entry.

    This function takes a dictionary containing registration and invoice data,
    formats it into a pandas DataFrame, and appends it to an Excel-based
    sponsor database. If the database file does not exist, it creates a new one
    with an appropriate Excel table structure. If it does exist, the function
    appends the new data and rewrites the file.

    :param registrer_dict: A dictionary containing all invoice-related fields,
        including customer information and products purchased.
    :type registrer_dict: dict[str, Any]
    :return: None
    :rtype: None
    """

    print("\t\t> Update invoice database...")
    
    # Sponsor database

    # Format sponsor DataFrame
    column_list = ["Date", "Invoice Number", "Company", "Title", "First Name", "Last Name", "Address", "Postcode", "City", "Phone", "Email", "Default Product Dict", "Custom Product Dict", "Total Price [CHF]", "Comment"]
    registrer_entry_list = list(registrer_dict.values())

    registrer_entry_df = pd.DataFrame([registrer_entry_list], columns=column_list)
    
    # Add data to database
    if not os.path.exists(SPONSOR_DATABASE_PATH):
        # If the database Excel file does NOT exist, create it and write the very first DataFrame
        with pd.ExcelWriter(SPONSOR_DATABASE_PATH, engine="xlsxwriter") as writer:
            # Write the sponsor entry data to the sheet
            registrer_entry_df.to_excel(writer, index=False, sheet_name=SHEET_NAME)
            # Getting XlsxWriter worksheet object
            worksheet = writer.sheets[SHEET_NAME]
            # Getting the dimensions of the DataFrame
            (max_row, max_col) = registrer_entry_df.shape
            # Creating a list of column headers, to use in "add_table()"
            column_settings = [{"header": column} for column in registrer_entry_df.columns]
            # Adding the Excel table structure (Pandas will add the data)
            worksheet.add_table(0, 0, max_row, max_col-1, {"columns": column_settings})
    else:
        # If the database Excel file exists, read the existing data, append the new data, and rewrite the file
        # Read the existing data
        existing_data_df = pd.read_excel(SPONSOR_DATABASE_PATH, sheet_name=SHEET_NAME)
        # Append the new data to the existing data
        combined_data_df = pd.concat([existing_data_df, registrer_entry_df], ignore_index=True)
        # Rewrite the combined data and recreate the pivot table
        with pd.ExcelWriter(SPONSOR_DATABASE_PATH, engine="xlsxwriter") as writer:
            # Write the combined data to the sheet
            combined_data_df.to_excel(writer, index=False, sheet_name=SHEET_NAME)
            # Getting XlsxWriter worksheet object
            worksheet = writer.sheets[SHEET_NAME]
            # Getting the dimensions of the DataFrame
            (max_row, max_col) = combined_data_df.shape
            # Creating a list of column headers, to use in "add_table()"
            column_settings = [{"header": column} for column in registrer_entry_df.columns]
            # Adding the Excel table structure (Pandas will add the data)
            worksheet.add_table(0, 0, max_row, max_col-1, {"columns": column_settings})

    print("\t\t\t> Invoice database successfully updated!")


def send_invoice_via_email(driver: WebDriver, registrer_dict: dict, invoice_path: str, index: int):
    """TODO: Description!!
    + Parameters!
    """
    print("\t\t> Send invoice via email...")

    # Click on button "Nouveau message"
    new_message_button_xpath = '//*[@id="step1"]'
    click_button(driver=driver, xpath=new_message_button_xpath)
    time.sleep(2)  # adjust if needed

    # Input recipient email address
    recipient_input_xpath = f'//*[@id="mat-chip-list-input-{str(index)}"]'
    #email = registrer_dict["email"]
    email = "antho.guinchard@gmail.com"
    enter_input(driver=driver, xpath=recipient_input_xpath, text_input=email)

    # Input email subject
    email_subject_input_xpath = f'//*[@id="mat-input-{str(8+index)}"]'
    invoice_name = Path(invoice_path).name
    invoice_number = re.search(r'\d+', invoice_name).group(0)
    email_subject = f"Giron du Nord 2025 à Concise • Facture inscription sports • {invoice_number}"
    enter_input(driver=driver, xpath=email_subject_input_xpath, text_input=email_subject)

    # Input email content
    editor_name = "squireEditor"
    email_content = compose_email(registrer_dict=registrer_dict)
    enter_editor(driver=driver, editor_name=editor_name, text_input=email_content)
    
    # Attach invoice
    # Wait for the file input element to be present
    file_input = WebDriverWait(driver, 10).until(
        EC.presence_of_element_located((By.XPATH, '//input[@type="file"]'))
    )
    # Send the file path to the input (this uploads the file)
    file_input.send_keys(invoice_path)
    time.sleep(2)  # adjust if needed

    # Send mail by clicking twice on button "ENVOYER" (to prevent pop-up message appearing in case of sending message during weekend)
    #send_message_button_xpath = '//*[@id="cdk-overlay-0"]/app-compose-dialog/div/div/div[2]/app-mail-composer/form/div[2]/div[2]/div/button[1]'
    send_message_button_xpath = "//button[normalize-space(.)='Envoyer']"
    click_button(driver=driver, xpath=send_message_button_xpath)

    # Wait a bit after sending email to make sure it has been sent
    time.sleep(5)  # adjust if needed

    try:
        recipient_input_element = driver.find_element(By.XPATH, recipient_input_xpath)
        print(colored("\t\t\tError!", "red"), "Recipient input element is still present. This means that email could not be sent...")
    except NoSuchElementException:
        print("\t\t\tRecipient input element no more present. Email successfully sent!")

@click.command()
@click.option("-d", "--debug", is_flag=True, help="Enable debug mode.", default=True)
def main(debug: bool):
    """Run script for generating and sending sport invoices.
    """
    # Set debug mode
    DEBUG_MODE = debug  # set DEBUG_MODE based on the flag
    
    if DEBUG_MODE:
        print("Debug mode is ON")
    else:
        print("Debug mode is OFF")

    # Set up print statement redirection to log file
    sys.stdout = DualLogger(LOG_PATH / f"{SCRIPT_NAME}_{CURRENT_TIME}.log")
    
    # Set up Selenium
    driver = setup_selenium()

    # Load sports catalog
    with open (SPORTS_CATALOG_PATH, "r") as file:
        global SPORTS_CATALOG_DICT
        SPORTS_CATALOG_DICT = json.load(file)

    # Read Excel file with sport registrations
    df = load_registrations_from_excel()

    # Sanitize data
    df_sanitized = sanitize_data(df)

    # Loop through each registration
    print("Processing registrations...")
    for index, row in df_sanitized.iterrows():
        print(f"\t{index + 1}/{len(df_sanitized)}: entry ID {row['Entry ID']} (name: {row['Name']};  email: {row['Email']})")
        
        # Generate invoice
        registrer_dict, invoice_path = generate_invoice(entry=row)

        # Update invoice database
        update_invoice_database(registrer_dict=registrer_dict)

        # Send invoice via email using selenium
        send_invoice_via_email(driver=driver, registrer_dict=registrer_dict, invoice_path=invoice_path, index=index)

    # Shut down Selenium
    shutdown_selenium(driver=driver)

    print("✅ Finished! All invoices have been generated and sent!")


if __name__ == "__main__":
    main()
