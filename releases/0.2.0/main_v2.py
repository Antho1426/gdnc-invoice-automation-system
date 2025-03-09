# Script name:         main.py
# Python interpreter:  Miniconda virtual environment "automation-env"
# Description:         Code for generating GDNC invoice to send to sponsor or donator
# Invocation example:  python main.py
# Author:              Anthony Guinchard
# Version:             0.1
# Creation date:       2024-09-04
# Modification date:   2025-03-09
# Resources:
#                      - YouTube video "NeuralNine - Invoice Automation System in Python - Full Project" (https://youtu.be/JuBEC1RW8nA?si=-a1BploFfwDJsV0a).
#                      - Discussion with ChatGPT (https://chatgpt.com/c/6792c47d-01b4-8003-abc4-23d175330cdc)
# Working:             ✅

from pathlib import Path

import datetime as dt
import json
import math as m
import os
import re
import subprocess
import time
import tkinter as tk
from datetime import datetime, timedelta
from enum import Enum
from functools import wraps
from tkinter import (BOTH, LEFT, RIGHT, VERTICAL, Canvas, Frame, Y, filedialog,
                     messagebox, ttk)

import docx
import pandas as pd
import requests
from docx2pdf import convert
from PIL import Image, ImageSequence, ImageTk

# ++++++++++++++++
DEBUG_MODE = True
# ++++++++++++++++
PROJECT_PATH = Path(os.getcwd())
SRC_PATH = PROJECT_PATH / "src"
ASSETS_PATH = SRC_PATH / "assets"
BIN_PATH = SRC_PATH / "bin"
LIB_PATH = SRC_PATH / "lib"


def validate_non_empty(cls):
    """Decorator to ensure no attribute is empty or an empty dictionary, with a warning popup instead of an error, except for special cases like Products where having one of default or custom product dict empty is allowed but not both."""
    
    orig_init = cls.__init__

    @wraps(orig_init)
    def new_init(self, *args, **kwargs):
        orig_init(self, *args, **kwargs)  # Call the original __init__

        missing_attrs = []
        
        # Special case for Products class
        if cls.__name__ == "Products":
            if not self.default and not self.custom:  # Both can't be empty
                missing_attrs.append("default and custom (at least one must be filled)")
        
        else:
            for attr, value in self.__dict__.items():
                if value in ("", {}, None):
                    missing_attrs.append(attr)
                    setattr(self, attr, "MISSING")  # Set a placeholder value

        if missing_attrs:
            messagebox.showwarning(
                title="Missing Information",
                message=f"The following fields are missing in {self.__class__.__name__} object:\n"
                        f"{', '.join(missing_attrs)}\n\n"
                        f"Please update them before proceeding."
            )

    cls.__init__ = new_init  # Reassign the wrapped function
    return cls


class SponsorObject:
    """Main class containing nested classes for structured sponsor data."""

    @validate_non_empty
    class Info:
        def __init__(self, company: str, title: str, first_name: str, last_name: str, address: str, postcode: str, city: str):
            self.company = company
            self.title = title
            self.first_name = first_name
            self.last_name = last_name
            self.address = address
            self.postcode = postcode
            self.city = city

    @validate_non_empty
    class Contact:
        def __init__(self, phone: str, email: str):
            self.phone = phone
            self.email = email

    @validate_non_empty
    class Invoice:
        def __init__(self, number: str, date: str, deadline: str):
            self.number = number
            self.date = date
            self.deadline = deadline

    @validate_non_empty
    class Products:
        def __init__(self, default: dict, custom: dict):
            self.default = default
            self.custom = custom
    
    def __init__(self, info, contact, invoice, products):
        self.info = info
        self.contact = contact
        self.invoice = invoice
        self.products = products

    def has_missing_values(self):
        """Check if any attribute (including nested ones) has the value 'MISSING'.
        Ensure at least one of default product dict or custom product dict is filled.
        """
        for section in [self.info, self.contact, self.invoice]:
            for attr, value in section.__dict__.items():
                if value == "MISSING":
                    return True  # found missing required values

        # Special validation for Products
        if not self.products.default and not self.products.custom:
            return True  # having both product dictionaries empty is not allowed

        return False  # everything is valid


class InvoiceAutomation:
    # Class-level constants
    PRODUCT_CATALOG_NAME = "product_catalog.json"  # source: "/Users/anthony/Dropbox/DocumentsPartagésMBPro↔MBAir/GironDuNord2025ÀConcise(GDNC)/Sponsoring/ContratSponsoring_v3.pdf"
    GDNC_LOGO_NAME = "gdnc.png"
    GDNC_LOGO_SPINNING_WHEEL_NAME = "gdnc-spinning-wheel.gif"
    GDNC_LOGO_CHECK_NAME = "gdnc-check.png"
    INVOICE_MODELS_FOLDER_NAME = "invoice_models"
    SPONSOR_DATABASE_NAME = "sponsor_database.xlsx"
    SHEET_NAME = "Sheet1"
    VERSION = "0.2.0"
    PAD = 5  # set a consistent padding for widgets
    SPINNING_IMAGE_SIZE = 70  # set spinning image size

    def __init__(self):
        # Create the main application window
        self.root = tk.Tk()
        self.root.title(f"GDNC Invoice Automation System {self.VERSION}")
        self.root.geometry("600x750")

        # --- Set up scrollbar for full window
        # Main frame
        main_frame = Frame(self.root)
        main_frame.pack(fill=BOTH, expand=1)
        # Canvas
        my_canvas = Canvas(main_frame)
        my_canvas.pack(side=LEFT, fill=BOTH, expand=1)
        # Scrollbar
        my_scrollbar = tk.Scrollbar(main_frame, orient=VERTICAL, command=my_canvas.yview)
        my_scrollbar.pack(side=RIGHT, fill=Y)
        # Configure the canvas
        my_canvas.configure(yscrollcommand=my_scrollbar.set)
        my_canvas.bind(
            '<Configure>', lambda e: my_canvas.configure(scrollregion=my_canvas.bbox("all"))
        )
        second_frame = Frame(my_canvas)
        # ---

        # ----------------------------------------------------------------------
        # ▷ Create the Info frame
        self.info_frame = ttk.LabelFrame(second_frame)
        self.info_frame.grid(row=0, column=0, padx=self.PAD, pady=self.PAD, sticky="ew")
        self.info_frame.configure(labelwidget=ttk.Label(self.info_frame, text="Info", font=("TkDefaultFont", 15, "bold")))

        ttkwidgets = ["Company", "Title", "Address", "Postcode", "City"]
        debug_value_info_list = ["Voyager SA", "Monsieur", "Via Ra Cürta 2", "6926", "Montagnola"]

        # Add Company row
        ttk.Label(self.info_frame, text=ttkwidgets[0]).grid(row=0, column=0, sticky="w", padx=self.PAD, pady=self.PAD)
        self.company = ttk.Entry(self.info_frame)
        if DEBUG_MODE:
            self.company.insert(0, debug_value_info_list[0])
        self.company.grid(row=0, column=1, columnspan=3, sticky="ew", padx=self.PAD, pady=self.PAD)

        # Add Title row
        ttk.Label(self.info_frame, text=ttkwidgets[1]).grid(row=1, column=0, sticky="w", padx=self.PAD, pady=self.PAD)
        self.title = ttk.Combobox(self.info_frame, values=["Monsieur", "Madame"], state="readonly")
        self.title.set("Monsieur")
        self.title.grid(row=1, column=1, columnspan=3, sticky="ew", padx=self.PAD, pady=self.PAD)
        
        # Add First Name and Last Name on the same horizontal line
        name_row = 2
        ttk.Label(self.info_frame, text="First Name").grid(row=name_row, column=0, sticky="w", padx=self.PAD, pady=self.PAD)
        self.first_name = ttk.Entry(self.info_frame)
        if DEBUG_MODE:
            self.first_name.insert(0, "Fabian")
        self.first_name.grid(row=name_row, column=1, sticky="ew", padx=self.PAD, pady=self.PAD)
        ttk.Label(self.info_frame, text="Last Name").grid(row=name_row, column=2, sticky="w", padx=self.PAD, pady=self.PAD)
        self.last_name = ttk.Entry(self.info_frame)
        if DEBUG_MODE:
            self.last_name.insert(0, "Giger")
        self.last_name.grid(row=name_row, column=3, sticky="ew", padx=self.PAD, pady=self.PAD)

        # Add Address, Postcode, and City rows
        for i, label in enumerate(ttkwidgets[2:], start=3):
            ttk.Label(self.info_frame, text=label).grid(row=i, column=0, sticky="w", padx=self.PAD, pady=self.PAD)
            #ttk.Entry(self.info_frame).grid(row=i, column=1, columnspan=3, sticky="ew", padx=self.PAD, pady=self.PAD)
            entry_widget = ttk.Entry(self.info_frame)
            entry_widget.insert(0, debug_value_info_list[i-1])
            entry_widget.grid(row=i, column=1, columnspan=3, sticky="ew", padx=self.PAD, pady=self.PAD)
            # Dynamically create attributes like self.address, self.postcode, etc.
            setattr(self, label.lower().strip(), entry_widget)

        self.info_frame.columnconfigure(1, weight=1)
        self.info_frame.columnconfigure(3, weight=1)

        # ----------------------------------------------------------------------
        # ▷ Create the Contact frame
        self.contact_frame = ttk.LabelFrame(second_frame)
        self.contact_frame.grid(row=1, column=0, padx=self.PAD, pady=self.PAD, sticky="ew")
        self.contact_frame.configure(labelwidget=ttk.Label(self.contact_frame, text="Contact", font=("TkDefaultFont", 15, "bold")))

        ttkwidgets_contact = ["Phone       ", "Email"]
        debug_value_contact_list = ["+41 79 123 45 67", "fabian.giger@gmail.com"]
        for i, label in enumerate(ttkwidgets_contact):
            ttk.Label(self.contact_frame, text=label).grid(row=i, column=0, sticky="w", padx=self.PAD, pady=self.PAD)
            #ttk.Entry(self.contact_frame).grid(row=i, column=1, sticky="ew", padx=self.PAD, pady=self.PAD)
            entry_widget = ttk.Entry(self.contact_frame)
            entry_widget.insert(0, debug_value_contact_list[i])
            entry_widget.grid(row=i, column=1, sticky="ew", padx=self.PAD, pady=self.PAD)
            # Dynamically create attributes like self.phone and self.email
            setattr(self, label.lower().strip(), entry_widget)
        self.contact_frame.columnconfigure(1, weight=1)

        # ----------------------------------------------------------------------
        # ▷ Create the Invoice frame
        self.invoice_frame = ttk.LabelFrame(second_frame)
        self.invoice_frame.grid(row=2, column=0, padx=self.PAD, pady=self.PAD, sticky="ew")
        self.invoice_frame.configure(labelwidget=ttk.Label(self.invoice_frame, text="Invoice", font=("TkDefaultFont", 15, "bold")))

        ttkwidgets_invoice = ["Number", "Date", "Deadline   "]
        ttkwidgets_invoice_values = [get_latest_invoice_number(), get_tomorrow_formatted_date(), get_deadline_formatted_date()]
        for i, label in enumerate(ttkwidgets_invoice):
            ttk.Label(self.invoice_frame, text=label).grid(row=i, column=0, sticky="w", padx=self.PAD, pady=self.PAD)
            #ttk.Entry(self.invoice_frame).grid(row=i, column=1, sticky="ew", padx=self.PAD, pady=self.PAD)
            entry_widget = ttk.Entry(self.invoice_frame)
            entry_widget.grid(row=i, column=1, sticky="ew", padx=self.PAD, pady=self.PAD)
            entry_widget.insert(0, ttkwidgets_invoice_values[i])
            # Dynamically create attributes like self.number, self.date, etc.
            setattr(self, label.lower().strip(), entry_widget)
        self.invoice_frame.columnconfigure(1, weight=1)

        # ----------------------------------------------------------------------
        # ▷ Create the Products frame
        self.products_frame = ttk.LabelFrame(second_frame)
        self.products_frame.grid(row=3, column=0, padx=self.PAD, pady=self.PAD, sticky="ew")
        self.products_frame.configure(labelwidget=ttk.Label(self.products_frame, text="Products", font=("TkDefaultFont", 15, "bold")))

        # Initialize lists to store selected default and custom product data (sub-dictionary names correspond to immutable "product creation order ID")
        self.selected_default_product_dict = {}
        self.selected_custom_product_dict = {}

        # Default product row

        with open (LIB_PATH / self.PRODUCT_CATALOG_NAME, "r") as file:
            self.default_product_catalog_dict = json.load(file)
        self.default_product_catalog_name_list = [item["name"] for item in self.default_product_catalog_dict.values()]
        self.default_product_catalog_price_list = [item["price"] for item in self.default_product_catalog_dict.values()]

        self.frame_default = ttk.Frame(self.products_frame)
        self.frame_default.grid(row=1, column=0, sticky="ew", pady=self.PAD)
        self.frame_default.grid_remove()  # hide default product frame by default

        # Checkbox
        self.show_default_products_var = tk.BooleanVar(value=False)  # checkbox to toggle default products frame defaults to invisible
        self.show_default_products_checkbox = ttk.Checkbutton(
            self.products_frame, text="Add Default Products", 
            variable=self.show_default_products_var, 
            command=self.toggle_default_products
        )
        self.show_default_products_checkbox.grid(row=0, column=0, sticky="w", padx=self.PAD, pady=self.PAD)

        self.default_label = ttk.Label(self.frame_default, text="Default   ")
        self.default_label.grid(row=0, column=0, sticky="w", padx=self.PAD)
        self.default_quantity_var = tk.StringVar(value="1")  # initialize with default value
        self.default_quantity_var.trace_add("write", self.update_selected_default_product_data)  # listen for changes
        self.default_quantity = ttk.Spinbox(self.frame_default, from_=1, to=10, width=2, textvariable=self.default_quantity_var, state="readonly")
        self.default_quantity.grid(row=0, column=1, padx=self.PAD)
        self.default_product = ttk.Combobox(self.frame_default, values=self.default_product_catalog_name_list)
        self.default_product.set(self.default_product_catalog_name_list[6])  # default product is by default product number "7" in "default_product_catalog_dict"
        self.default_product.bind("<<ComboboxSelected>>", self.update_selected_default_product_data)  # # listen for changes, update the price dynamically when a user selects a product from the combobox
        self.default_product.grid(row=0, column=2, padx=self.PAD, sticky="ew")
        self.default_price = ttk.Entry(self.frame_default, width=10)
        self.selected_product_price = self.get_product_price(self.default_product.get())
        self.default_price.insert(0, f"{self.selected_product_price} CHF")
        self.default_price.config(state="readonly")  # make the price field non-editable
        self.default_price.grid(row=0, column=3, padx=self.PAD)
        ttk.Button(self.frame_default, text="➕", command=self.add_default_product_row).grid(row=0, column=4, padx=self.PAD)

        # Custom product row
        
        self.frame_custom = ttk.Frame(self.products_frame)
        self.frame_custom.grid(row=3, column=0, sticky="ew", pady=self.PAD)
        self.frame_custom.grid_remove()  # hide custom product frame by default

        # Checkbox
        self.show_custom_products_var = tk.BooleanVar(value=False)  # checkbox to toggle custom products frame defaults to invisible
        self.show_custom_products_checkbox = ttk.Checkbutton(
            self.products_frame, text="Add Custom Products", 
            variable=self.show_custom_products_var, 
            command=self.toggle_custom_products
        )
        self.show_custom_products_checkbox.grid(row=2, column=0, sticky="w", padx=self.PAD, pady=self.PAD)

        self.custom_label = ttk.Label(self.frame_custom, text="Custom  ")
        self.custom_label.grid(row=0, column=0, sticky="w", padx=self.PAD)
        self.custom_quantity_var = tk.StringVar(value="1")  # initialize with default value
        self.custom_quantity_var.trace_add("write", self.update_selected_custom_product_data)  # listen for changes
        self.custom_quantity = ttk.Spinbox(self.frame_custom, from_=1, to=10, width=2, textvariable=self.custom_quantity_var, state="readonly")
        self.custom_quantity.grid(row=0, column=1, padx=self.PAD)
        self.custom_product_var = tk.StringVar(value="Donation selon contrat")  # initialize with default value
        self.custom_product_var.trace_add("write", self.update_selected_custom_product_data)  # listen for changes
        self.custom_product = ttk.Entry(self.frame_custom, width=21, textvariable=self.custom_product_var)
        self.custom_product.grid(row=0, column=2, padx=self.PAD, sticky="ew")
        self.custom_price_var = tk.StringVar(value="100")  # initialize with default value
        self.custom_price_var.trace_add("write", self.update_selected_custom_product_data)  # listen for changes
        self.custom_price = ttk.Entry(self.frame_custom, width=5, textvariable=self.custom_price_var)
        self.custom_price.grid(row=0, column=3, padx=self.PAD)
        self.custom_currency_label = ttk.Label(self.frame_custom, text="CHF")
        self.custom_currency_label.grid(row=0, column=4, sticky="w", padx=self.PAD)
        ttk.Button(self.frame_custom, text="➕", command=self.add_custom_product_row).grid(row=0, column=5, padx=self.PAD)

        self.products_frame.columnconfigure(0, weight=1)

        # ----------------------------------------------------------------------
        # ▷ Price frame
        self.price_frame = ttk.Frame(second_frame)
        self.price_frame.grid(row=4, column=0, padx=self.PAD, pady=self.PAD, sticky="w")

        ttk.Label(self.price_frame, text="   Price", font=("TkDefaultFont", 15, "bold")).grid(row=0, column=0, sticky="e", padx=self.PAD)
        self.price_entry = ttk.Entry(self.price_frame, width=10)
        self.price_entry.insert(0, "0 CHF")
        self.price_entry.config(state="readonly")
        self.price_entry.grid(row=0, column=1, sticky="w", padx=self.PAD)
        
        if DEBUG_MODE:
            # Toggle default product checkbox to show and select 1 product
            self.show_default_products_var.set(True)
            self.toggle_default_products()  # manually call the function linked to the checkbox
             # Toggle custom product checkbox to show and select 1 product
            self.show_custom_products_var.set(True)
            self.toggle_custom_products()  # manually call the function linked to the checkbox
            
        # Spinning wheel
        # Static image
        self.static_img_path = ASSETS_PATH / self.GDNC_LOGO_NAME
        self.static_img = Image.open(self.static_img_path)
        self.static_img = self.static_img.resize((self.SPINNING_IMAGE_SIZE, self.SPINNING_IMAGE_SIZE), Image.LANCZOS)
        # Static image success
        self.static_img_success_path = ASSETS_PATH / self.GDNC_LOGO_CHECK_NAME
        self.static_img_success = Image.open(self.static_img_success_path)
        self.static_img_success = self.static_img_success.resize((self.SPINNING_IMAGE_SIZE, self.SPINNING_IMAGE_SIZE), Image.LANCZOS)
        self.success = False
        # GIF
        self.gif_path = ASSETS_PATH / self.GDNC_LOGO_SPINNING_WHEEL_NAME
        self.gif = Image.open(self.gif_path)
        self.frames = [ImageTk.PhotoImage(frame) for frame in ImageSequence.Iterator(self.gif)]
        self.frame_count = len(self.frames)
        self.current_frame = 0
        # Label for displaying the static image or GIF
        self.gif_label = tk.Label(self.price_frame)
        self.gif_label.grid(row=0, column=2, sticky="ew", padx=self.PAD*12, pady=self.PAD)
        # Disable the GIF by default (only show static image)
        self.start_spinning = True  # control GIF animation state
        self.toggle_spinning()
 
        # Multi-line status label
        self.status_label = ttk.Label(
            self.price_frame, 
            text="",
            font=("Courier", 10),
            justify="left",
            wraplength=180  # adjust this value to control line wrapping
        )
        self.status_label.grid(row=0, column=3, sticky="w", padx=self.PAD)

        # ----------------------------------------------------------------------
        # ▷ Create Invoice button
        self.create_invoice_button = ttk.Button(second_frame, text="Create invoice", command=self.create_invoice)
        self.create_invoice_button.grid(row=5, column=0, padx=self.PAD, pady=self.PAD, sticky="ew")

        # --- Create window around the target frame "second_frame" to finish setting up the scrollbar
        my_canvas.create_window((0, 0), window=second_frame, anchor="nw")
        # ---

        self.root.mainloop()
    

    def toggle_default_products(self):
        if self.show_default_products_var.get():
            self.frame_default.grid(row=1, column=0, sticky="ew", pady=self.PAD)
            # Iterate over frame_default widgets to gather data from the default product rows and rehabilitate selected_default_product_dict as it was before the frame was hidden
            row_count = 0
            for widget in self.frame_default.winfo_children():
                if isinstance(widget, ttk.Spinbox):
                    default_quantity = widget.get()
                if isinstance(widget, ttk.Combobox):
                    default_product = widget.get()
                if isinstance(widget, ttk.Button):
                    # Store reference of the products
                    self.selected_default_product_dict[f"{row_count}"] = {
                        "name": default_product,
                        "quantity": default_quantity
                    }
                    # Update row_count to begin gathering data from the next product row
                    row_count = row_count+1
        else:
            self.frame_default.grid_remove()
            self.selected_default_product_dict.clear()  # completely empty dictionary of selected default products
        self.print_selected_product_summary()
        # Update total price
        self.update_total_price()


    def update_selected_default_product_data(self, event, *args):
        selected_product = self.default_product.get()  # get selected product name
        if selected_product in self.default_product_catalog_name_list:
            # Update displayed price
            selected_price = self.get_product_price(selected_product)
            self.default_price.config(state="normal")  # enable editing
            self.default_price.delete(0, tk.END)  # clear the field
            self.default_price.insert(0, f"{selected_price} CHF")  # insert updated price
            self.default_price.config(state="readonly")  # disable editing again
            # Update dictionary of selected default products
            self.selected_default_product_dict["0"] = {  # update reference of the first product
                "name": selected_product,
                "quantity": self.default_quantity_var.get(),
            }
            self.print_selected_product_summary()
            # Update total price
            self.update_total_price()
        else:
            messagebox.showerror(title="Error", message="Non-existing default product selected!")
            return


    def add_default_product_row(self):
        """Dynamically adds a new default product row."""
        
        def update_new_selected_default_product_data(event, *args):
            selected_product = new_product.get()  # get selected product name
            if selected_product in self.default_product_catalog_name_list:
                # Update displayed price
                selected_price = self.get_product_price(selected_product)
                new_price.config(state="normal")  # enable editing
                new_price.delete(0, tk.END)  # clear the field
                new_price.insert(0, f"{selected_price} CHF")  # insert updated price
                new_price.config(state="readonly")  # disable editing again
                # Update dictionary of selected default products (i.e., update reference of the "row"-th product)
                self.selected_default_product_dict[f"{row}"] = {
                    "name": selected_product,
                    "quantity": new_quantity_var.get(),
                }
                self.print_selected_product_summary()
                # Update total price
                self.update_total_price()
            else:
                messagebox.showerror(title="Error", message="Non-existing default product selected!")
                return

        def remove_default_product_row():
            new_quantity.destroy()
            new_product.destroy()
            new_price.destroy()
            remove_button.destroy()
            self.selected_default_product_dict.pop(f"{row}")
            self.print_selected_product_summary()
            # Update total price
            self.update_total_price()

        # The "row" (i.e., current "product ID" in selected_default_product_dict)
        # is determined based on the number of already stored products
        row = len(self.selected_default_product_dict)

        # Quantity
        new_quantity_var = tk.StringVar(value="1")  # initialize with default value
        new_quantity_var.trace_add("write", update_new_selected_default_product_data)  # listen for quantity changes
        new_quantity = ttk.Spinbox(self.frame_default, from_=1, to=10, width=2, textvariable=new_quantity_var, state="readonly")
        new_quantity.grid(row=row, column=1, padx=self.PAD)

        # Product Name
        new_product = ttk.Combobox(self.frame_default, values=self.default_product_catalog_name_list)
        new_product.set(self.default_product_catalog_name_list[6])  # default product is by default product number "7" in "default_product_catalog_dict"
        new_product.bind("<<ComboboxSelected>>", update_new_selected_default_product_data)  # update the price dynamically when a user selects a product from the combobox
        new_product.grid(row=row, column=2, padx=self.PAD, sticky="ew")
        
        # Price
        new_price = ttk.Entry(self.frame_default, width=10)
        selected_new_price = self.get_product_price(new_product.get())
        new_price.insert(0, f"{selected_new_price} CHF")
        new_price.config(state="readonly")  # make the price field non-editable
        new_price.grid(row=row, column=3, padx=self.PAD)

        # Remove button
        remove_button = ttk.Button(self.frame_default, text="❌", command=remove_default_product_row)
        remove_button.grid(row=row, column=4, padx=self.PAD)

        # Update selected_default_product_dict
        self.selected_default_product_dict[f"{row}"] = {
            "name": new_product.get(),
            "quantity": new_quantity.get(),
        }

        self.print_selected_product_summary()

        # Update total price
        self.update_total_price()


    def toggle_custom_products(self):
        if self.show_custom_products_var.get():
            self.frame_custom.grid(row=3, column=0, sticky="ew", pady=self.PAD)
            # Iterate over frame_custom widgets to gather data from the custom product rows and rehabilitate selected_custom_product_dict as it was before the frame was hidden
            row_count = 0
            first_entry_not_yet_encountered = True
            for widget in self.frame_custom.winfo_children():
                if isinstance(widget, ttk.Spinbox):
                    custom_quantity = widget.get()
                if not isinstance(widget, ttk.Spinbox) and isinstance(widget, ttk.Entry) and first_entry_not_yet_encountered:
                    custom_product = widget.get()
                    first_entry_not_yet_encountered = False
                    continue
                if not isinstance(widget, ttk.Spinbox) and isinstance(widget, ttk.Entry) and not first_entry_not_yet_encountered:
                    custom_price = widget.get()
                if isinstance(widget, ttk.Button):
                    # Store reference of the products
                    self.selected_custom_product_dict[f"{row_count}"] = {
                        "name": custom_product,
                        "quantity": custom_quantity,
                        "price": custom_price
                    }
                    # Update row_count to begin gathering data from the next product row
                    row_count = row_count+1
        else:
            self.frame_custom.grid_remove()
            self.selected_custom_product_dict.clear()  # completely empty dictionary of selected custom products
        self.print_selected_product_summary()
        # Update total price
        self.update_total_price()
        

    def update_selected_custom_product_data(self, event, *args):
        # Prevent user from entering non-numeric characters
        self.custom_price_var.set("".join(filter(str.isdigit, self.custom_price_var.get())))
        custom_price_selected = self.custom_price_var.get()
        # Handle empty price field
        if self.custom_price_var.get() == "":
            custom_price_selected = "0"    
        # Update dictionary of selected custom products
        self.selected_custom_product_dict["0"] = {
            "name": self.custom_product_var.get(),
            "quantity": self.custom_quantity_var.get(),
            "price": custom_price_selected
        }
        self.print_selected_product_summary()
        # Update total price
        self.update_total_price()
        

    def add_custom_product_row(self):
        """Dynamically adds a new custom product row."""

        def update_new_selected_custom_product_data(event=None, *args):
            # Prevent user from entering non-numeric characters
            new_custom_price_var.set("".join(filter(str.isdigit, new_custom_price_var.get())))
            new_custom_price_selected = new_custom_price_var.get()
            # Handle empty price field
            if new_custom_price_var.get() == "":
                new_custom_price_selected = "0"    
            # Update dictionary of selected custom products
            self.selected_custom_product_dict[f"{row}"] = {
                "name": new_custom_product_var.get(),
                "quantity": new_custom_quantity_var.get(),
                "price": new_custom_price_selected
            }
            self.print_selected_product_summary()
            # Update total price
            self.update_total_price()

        def remove_custom_product_row():
            new_custom_quantity.destroy()
            new_custom_product.destroy()
            new_custom_price.destroy()
            new_custom_currency_label.destroy()
            remove_button.destroy()
            self.selected_custom_product_dict.pop(f"{row}")
            # Update total price
            self.update_total_price()
        
        # The "row" (i.e., current "product ID" in selected_custom_product_dict)
        # is determined based on the number of already stored products
        row = len(self.selected_custom_product_dict)

        # Quantity
        new_custom_quantity_var = tk.StringVar(value="1")  # initialize with default value
        new_custom_quantity_var.trace_add("write", update_new_selected_custom_product_data)  # listen for changes
        new_custom_quantity = ttk.Spinbox(self.frame_custom, from_=1, to=10, width=2, textvariable=new_custom_quantity_var, state="readonly")
        new_custom_quantity.grid(row=row, column=1, padx=self.PAD)

        # Product Name
        new_custom_product_var = tk.StringVar(value="Donation selon contrat")  # initialize with default value
        new_custom_product_var.trace_add("write", update_new_selected_custom_product_data)  # listen for changes
        new_custom_product = ttk.Entry(self.frame_custom, width=21, textvariable=new_custom_product_var)
        new_custom_product.grid(row=row, column=2, padx=self.PAD, sticky="ew")
        
        # Price
        new_custom_price_var = tk.StringVar(value="100")  # initialize with default value
        new_custom_price_var.trace_add("write", update_new_selected_custom_product_data)  # listen for changes
        new_custom_price = ttk.Entry(self.frame_custom, width=5, textvariable=new_custom_price_var)
        new_custom_price.grid(row=row, column=3, padx=self.PAD)
        
        # Currency label
        new_custom_currency_label = ttk.Label(self.frame_custom, text="CHF")
        new_custom_currency_label.grid(row=row, column=4, sticky="w", padx=self.PAD)

        # Remove button
        remove_button = ttk.Button(self.frame_custom, text="❌", command=remove_custom_product_row)
        remove_button.grid(row=row, column=5, padx=self.PAD)

        # Update selected_custom_product_dict
        self.selected_custom_product_dict[f"{row}"] = {
            "name": new_custom_product.get(),
            "quantity": new_custom_quantity.get(),
            "price": new_custom_price.get(),
        }

        self.print_selected_product_summary()

        # Update total price
        self.update_total_price()


    def get_product_price(self, selected_product: str):
        price = self.default_product_catalog_price_list[
                self.default_product_catalog_name_list.index(selected_product)
            ]
        return price


    def compute_total_price(self):
        """Compute total price of selected products."""
        total_price_default_products = 0
        for product in self.selected_default_product_dict.values():
            product_quantity = int(product["quantity"])
            product_price = self.get_product_price(product["name"])
            total_price_default_products = total_price_default_products + product_quantity * product_price
        
        total_price_custom_products = 0
        for product in self.selected_custom_product_dict.values():
            product_quantity = int(product["quantity"])
            product_price = int(product["price"])
            total_price_custom_products = total_price_custom_products + product_quantity * product_price
        
        total_price = total_price_default_products + total_price_custom_products

        return total_price


    def update_total_price(self) -> None:
        self.total_price = self.compute_total_price()
        self.price_entry.config(state="normal")  # enable editing
        self.price_entry.delete(0, tk.END)
        self.price_entry.insert(0, f"{self.total_price} CHF")
        self.price_entry.config(state="readonly")


    def print_selected_product_summary(self) -> None:
        """Retrieves selected product and print data."""
        print("\n---\nSelected default products:\n", json.dumps(self.selected_default_product_dict, indent=4, ensure_ascii=False))
        print("Selected custom products:\n", json.dumps(self.selected_custom_product_dict, indent=4, ensure_ascii=False))


    @staticmethod
    def replace_text(paragraph, old_text, new_text):
        """Replaces old_text with new_text in the given paragraph while keeping the original font and bold style.
        """
        if old_text in paragraph.text:
            bold_text_list = ["[TOTAL]", "[COMPANY]"]
            bold = False
            # Capture the fact that text has to be bold for the "TOTAL" text
            if any(keyword in paragraph.text for keyword in bold_text_list):
                bold = True
            # Replace the text
            paragraph.text = paragraph.text.replace(old_text, new_text)
            # Apply font style
            paragraph.style.font.name = "Times New Roman"
            # Apply bold
            if bold:
                paragraph.style.font.bold = True
            else:
                paragraph.style.font.bold = False
                for run in paragraph.runs:
                    run.bold = False
            if DEBUG_MODE:
                print(f"---\nparagraph.text: {paragraph.text}")
                print(f"paragraph.style.font.bold: {paragraph.style.font.bold}")


    def animate_gif(self):
        """Animate a transparent GIF properly"""
        if self.start_spinning:
            frame = self.gif.seek(self.current_frame)  # move to next frame

            # Convert to RGBA to handle transparency
            frame = self.gif.convert("RGBA")

            # Resize the frame dynamically (adjust width & height as needed)
            new_size = (self.SPINNING_IMAGE_SIZE, self.SPINNING_IMAGE_SIZE)
            frame = frame.resize(new_size, Image.LANCZOS)

            # Create a transparent-friendly background
            background = Image.new("RGBA", frame.size, (255, 255, 255, 0))  # white background
            frame = Image.alpha_composite(background, frame)  # blend the GIF frame
            
            # Convert to a Tkinter-compatible format
            self.frames[self.current_frame] = ImageTk.PhotoImage(frame)
            self.gif_label.config(image=self.frames[self.current_frame])

            # Cycle through frames
            self.current_frame = (self.current_frame + 1) % self.frame_count

            # Call again to continue animation
            self.root.after(self.gif.info['duration'], self.animate_gif)
        else:
            self.show_static_image()


    def show_static_image(self):
        """Switch to a static image when animation stops"""
        if self.success:
            static_img_success = ImageTk.PhotoImage(self.static_img_success)
            self.gif_label.config(image=static_img_success)
            self.gif_label.image = static_img_success  # keep reference to prevent garbage collection
        else:
            static_img = ImageTk.PhotoImage(self.static_img)
            self.gif_label.config(image=static_img)
            self.gif_label.image = static_img  # keep reference to prevent garbage collection
            self.success = False  # reset success variable


    def toggle_spinning(self):
        """Toggle spinning based on self.start_spinning"""
        self.start_spinning = not self.start_spinning
        if self.start_spinning:
            self.animate_gif()  # start animation
        else:
            self.show_static_image()  # stop animation


    def create_invoice(self):
        """Create invoice in PDF format.
        TODO: Update docstring here!
        """

        self.print_selected_product_summary()
        
        # Check that we have at least 1 and maximum 5 products (default + custom) selected
        num_selected_default_products = len(self.selected_default_product_dict)
        num_selected_custom_products = len(self.selected_custom_product_dict)
        if num_selected_default_products == 0 and num_selected_custom_product == 0:
            messagebox.showerror(title="Error", message="No product selected!")
            return
        num_total_products = num_selected_default_products + num_selected_custom_products
        if num_total_products > 5:
            messagebox.showerror(title="Error", message=f"⚠️ Maximum 5 products allowed!\nNumber of products currently selected:\n- Default: {str(num_selected_default_products)}\n- Custom: {str(num_selected_custom_products)}")
            return

        # Create SponsorObject object with attribute values retrieved from values entered by user
        sponsor = SponsorObject(
            info=SponsorObject.Info(
                company=self.company.get(),
                title=self.title.get(),
                first_name=self.first_name.get(),
                last_name=self.last_name.get(),
                address=self.address.get(),
                postcode=self.postcode.get(),
                city=self.city.get(),
            ),
            contact=SponsorObject.Contact(
                phone=self.phone.get(),
                email=self.email.get(),
            ),
            invoice=SponsorObject.Invoice(
                number=self.number.get(),
                date=self.date.get(),
                deadline=self.deadline.get(),
            ),
            products=SponsorObject.Products(
                default=self.selected_default_product_dict,
                custom=self.selected_custom_product_dict,
            )
        )

        # Check that none of the fields are left empty (i.e., that our SponsorObject object has values for all attributes) and stop execution if any value is "MISSING" (not: At this stae, warning messages have already been displayed to the user by decorator of subclasses of class SponsorObject when creating the sponsor object above)
        if sponsor.has_missing_values():
            # TODO: Update below as log message!
            print("⚠️ SponsorObject contains missing values. Please update them.")
            return  # stops further execution

        # Toggle spinning animation indicating that app is running
        self.toggle_spinning()
        
        # Update status label
        status_text = "Process launched! 🚀\n(👀 See terminal for outputs)"
        self.status_label.config(text=status_text)
        self.root.update()

        # Generate DOCX document

        # Update status label
        status_text = status_text + "\n> Generate DOCX invoice..."
        self.status_label.config(text=status_text)
        self.root.update()

        template_path = LIB_PATH / self.INVOICE_MODELS_FOLDER_NAME / f"InvoiceModel_CH95_DefaultProducts_{num_total_products}.docx"
        doc = docx.Document(template_path)

        product_key = "[PRODUCT-DESCRIPTION-IDX]"
        quantity_key = "[QT-IDX]"
        price_key = "[P-IDX]"
        tot_key = "[TOT-IDX]"

        replacements = {
            "[COMPANY]": sponsor.info.company,
            "[TITLE]": sponsor.info.title,
            "[FIRST-NAME]": sponsor.info.first_name,
            "[LAST-NAME]": sponsor.info.last_name,
            "[ADDRESS]": sponsor.info.address,
            "[POSTCODE]": sponsor.info.postcode,
            "[CITY]": sponsor.info.city,
            "[INVOICE-NUMBER]": sponsor.invoice.number,
            "[ISSUE-DATE]": sponsor.invoice.date,
            "[DEADLINE-DATE]": sponsor.invoice.deadline,
            "[TOTAL]": str("{:.2f}".format(float(self.total_price)))  # convert price to string since all mapped values have to have type string
        }

        default_product_replacements = {}
        for i in range(num_selected_default_products):
            name = sponsor.products.default[str(i)]["name"]
            quantity = sponsor.products.default[str(i)]["quantity"]
            price = "{:.2f}".format(float(self.get_product_price(name)))
            idx = i+1
            product_key_idx = product_key.replace("IDX", str(idx))
            default_product_replacements[product_key_idx] = name
            quantity_key_idx = quantity_key.replace("IDX", str(idx))
            default_product_replacements[quantity_key_idx] = quantity
            price_key_idx = price_key.replace("IDX", str(idx))
            default_product_replacements[price_key_idx] = str(price)  # convert price to string since all mapped values have to have type string
            tot_key_idx = tot_key.replace("IDX", str(idx))
            default_product_replacements[tot_key_idx] = str("{:.2f}".format(int(quantity)*float(price)))

        replacements.update(default_product_replacements)

        custom_product_replacements = {}
        for i in range(num_selected_custom_products):
            name = sponsor.products.custom[str(i)]["name"]
            quantity = sponsor.products.custom[str(i)]["quantity"]
            price = str("{:.2f}".format(float(sponsor.products.custom[str(i)]["price"])))
            idx = i+1+num_selected_default_products  # account for default products already registered in invoice by adding num_selected_default_products to idx of custom products
            product_key_idx = product_key.replace("IDX", str(idx))
            custom_product_replacements[product_key_idx] = name
            quantity_key_idx = quantity_key.replace("IDX", str(idx))
            custom_product_replacements[quantity_key_idx] = quantity
            price_key_idx = price_key.replace("IDX", str(idx))
            custom_product_replacements[price_key_idx] = price
            tot_key_idx = tot_key.replace("IDX", str(idx))
            custom_product_replacements[tot_key_idx] = str("{:.2f}".format(int(quantity)*float(price)))

        replacements.update(custom_product_replacements)
        
        # Make replacements in the paragraphs of the DOCX document (i.e., info and invoice data)
        for paragraph in list(doc.paragraphs):
            for old_text, new_text in replacements.items():
                self.replace_text(paragraph=paragraph,
                                  old_text=old_text, new_text=new_text)

        # Make replacements in the tables of the DOCX document (i.e., product data)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            self.replace_text(
                                paragraph=paragraph, old_text=old_text, new_text=new_text)
                        # Make sure product number "01", "02", etc. (under the column "NO" in the invoice template) are not bold
                        if re.fullmatch(r"0[1-5]", paragraph.text):
                            paragraph.style.font.bold = False
                            for run in paragraph.runs:
                                run.bold = False

        output_docx_path = f"invoice_populated/Facture N° {sponsor.invoice.number}.docx"
        doc.save(output_docx_path)
        
        # Convert DOCX to PDF

        #output_pdf_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF documents", "*.pdf")])
        output_pdf_path = output_docx_path.replace(".docx", ".pdf")

        # Notes:
        # - For converting DOCX to PDF, Microsoft Word has to be installed and already opened! In addition, access to the folder in which the generated invoice will be saved has first to be granted to Microsoft Word via the Microsoft pop up message!
        # - An internet connection is required to convert DOCX to PDF
        
        # Open up Microsoft Word to save time for future invoice generations
        status_text = status_text + "\n> Open up Microsoft Word..."
        self.status_label.config(text=status_text)
        self.root.update()

        # TODO: Programmatically open Microsoft Word if not yet opened!

        # Update status label
        status_text = status_text + "\n> Check internet connection..."
        self.status_label.config(text=status_text)
        self.root.update()

        # Check internet connection
        if check_internet():
            print("Internet is available!")
        else:
            messagebox.showerror(title="Error", message=f"No internet connection. The DOCX invoice could be generated but not converted into PDF. Invoice generation will stop here.")
            return

        # Update status label
        status_text = status_text + "\n> Convert DOCX to PDF..."
        self.status_label.config(text=status_text)
        self.root.update()
        
        convert(input_path=output_docx_path, output_path=output_pdf_path)

        # Compose email to send

        # Update status label
        status_text = status_text + "\n> Compose email to send..."
        self.status_label.config(text=status_text)
        self.root.update()

        # Sponsor email address
        print(f"\n▷ Sponsor email address:\n---\n{sponsor.contact.email}\n---")

        # Email subject
        subject = f"Facture sponsoring Giron du Nord 2025 à Concise • {sponsor.invoice.number}"
        print(f"\n▷ Generated email subject:\n---\n{subject}\n---")

        # Determine email to send based on current  day of the week and current hour
        
        current_hour = datetime.now().hour
        current_day = datetime.now().weekday()  # Monday = 0, Tuesday = 1, Wednesday = 2, Thursday = 3, Friday = 4, Saturday = 5, Sunday = 6

        email_template = "[GREETING] [TITLE] [LAST-NAME],\n\nSuite au contact que vous avez récemment eu avec notre organisation, nous vous envoyons la facture de sponsoring conformément au contrat que vous avez signé et que vous nous avez transmis.\n\nNous restons à votre entière disposition pour toute précision. Merci encore beaucoup de votre soutien dans cette aventure!\n\nMeilleures salutations et [CLOSING],\n\nAnthony Guinchard\nCommission Finances\nGiron du Nord 2025 à Concise"

        class Greeting(Enum):
            MORNING = "Bonjour"
            EVENING = "Bonsoir"

        if current_hour < 18:
            greeting = Greeting.MORNING.value
        else:
            greeting = Greeting.EVENING.value

        
        class Closing(Enum):
            START_OF_WEEK = "encore un bon début de semaine"
            MID_WEEK = "une bonne suite de semaine"
            END_OF_WEEK = "déjà une bonne fin de semaine"
            WEEKEND = "un bon weekend"
            END_OF_WEEKEND = "déjà une bonne fin de weekend"

        match current_day:
            case 0 | 1:  # Monday or Thursday
                closing = Closing.START_OF_WEEK.value
            case 2 | 3: # Wednesday or Tuesday
                closing = Closing.MID_WEEK.value
            case 4: # Friday
                closing = Closing.END_OF_WEEK.value
            case 5: # Saturday
                closing = Closing.WEEKEND.value
            case 6: # Sunday
                closing = closing.END_OF_WEEKEND.value
        
        email = email_template.replace("[GREETING]", greeting).replace("[TITLE]", sponsor.info.title).replace("[LAST-NAME]", sponsor.info.last_name).replace("[CLOSING]", closing)
        
        print(f"\n▷ Generated email:\n---\n{email}\n---")

        # Sponsor data backup and database update

        # Update status label
        status_text = status_text + "\n> Update sponsor database..."
        self.status_label.config(text=status_text)
        self.root.update()

        # Generate new line to manually copy-paste in file "1_N° facture.xlsx"

        today = datetime.today().strftime("%d.%m.%Y")
        numero_facture_entry = f"{today}\t{sponsor.invoice.number}\t{sponsor.info.company}\t{self.total_price:.2f} CHF\tMail"
        print(f"\n▷ Generated line for file '1_N° facture.xlsx':\n---\n{numero_facture_entry}\n---")

        # Sponsor database

        # TODO: Convert address into geographic coordinates and add such a column "Geographic Coordinates"

        # Format sponsor DataFrame
        column_list = ["Date", "Invoice Number", "Company", "Title", "First Name", "Last Name", "Address", "Postcode", "City", "Phone", "Email", "Default Product Dict", "Custom Product Dict", "Total Price [CHF]"]
        sponsor_entry_list = [today, sponsor.invoice.number, sponsor.info.company, sponsor.info.title, sponsor.info.first_name, sponsor.info.last_name, sponsor.info.address, sponsor.info.postcode, sponsor.info.city, sponsor.contact.phone, sponsor.contact.email, sponsor.products.default, sponsor.products.custom, self.total_price]
        sponsor_entry_df = pd.DataFrame([sponsor_entry_list], columns=column_list)

        sponsor_database_path = LIB_PATH / self.SPONSOR_DATABASE_NAME
        
        # Append data to database
        if not os.path.exists(sponsor_database_path):
            # If the database Excel file does NOT exist, create it and write the DataFrame
            with pd.ExcelWriter(sponsor_database_path) as writer:
                sponsor_entry_df.to_excel(writer, index=False, sheet_name=self.SHEET_NAME)
                # Getting XlsxWriter workbook and worksheet objects
                workbook = writer.book
                worksheet = writer.sheets[self.SHEET_NAME]
                # Getting the dimensions of the DataFrame
                (max_row, max_col) = sponsor_entry_df.shape
                # Creating a list of column headers, to use in "add_table()"
                column_settings = [{"header": column} for column in sponsor_entry_df.columns]
                # Adding the Excel table structure (Pandas will add the data)
                worksheet.add_table(0, 0, max_row, max_col-1, {"columns": column_settings})
        
        else:
            # TODO: Ask an LLM to modify below so that it doesn't just add line but is part of the excel pivot table (we should see the alternating blue and white colors!). Because those added lines are not part of the pivot table then!
            with pd.ExcelWriter(sponsor_database_path, engine="openpyxl", mode="a", if_sheet_exists="overlay") as writer:
                # Write the new data at the bottom of the sheet (append mode)
                sponsor_entry_df.to_excel(writer, sheet_name=self.SHEET_NAME, startrow=writer.sheets[self.SHEET_NAME].max_row, index=False, header=False)
        
        # Update status label
        status_text = status_text + "\nProcess finished! ✅"
        self.status_label.config(text=status_text)
        self.root.update()

        # Display success logo
        self.success = True
        self.toggle_spinning()

        # Display success message box
        messagebox.showinfo(
            title="Success", message="Invoice created and saved successfully!")


def get_tomorrow_date() -> datetime:
    today = datetime.now()
    tomorrow = today + timedelta(days=1)
    return tomorrow

def get_tomorrow_formatted_date() -> str:
    """Get date of tomorrow in formatted format "dd.mm.yyyy".
    """
    tomorrow_formatted = get_tomorrow_date().strftime("%d.%m.%Y")
    return tomorrow_formatted

def get_deadline_formatted_date() -> str:
    """Get date of tomorrow plus 30 days in formatted format "dd.mm.yyyy".
    """
    in_30_days = get_tomorrow_date() + timedelta(days=30)
    in_30_days_formatted = in_30_days.strftime("%d.%m.%Y")
    return in_30_days_formatted

def get_latest_invoice_number() -> str:
    """Retrieve latest invoice number from Excel file
    TODO: Update this function doctstring + real functionality!!
    """
    return 20240023


def check_internet(url="https://www.google.com", timeout=3):
    try:
        requests.get(url, timeout=timeout)
        return True
    except requests.RequestException:
        return False


if __name__ == "__main__":
    InvoiceAutomation()
