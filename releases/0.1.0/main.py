# Script name:         main.py
# Python interpreter:  Miniconda virtual environment "automation-env"
# Description:         Code for generating GDNC invoice to send to sponsor or donator
# Invocation example:  python main.py
# Author:              Anthony Guinchard
# Version:             0.1
# Creation date:       2024-09-04
# Modification date:   2024-10-17
# Resource:            YouTube video "NeuralNine - Invoice Automation System in Python - Full Project" (https://youtu.be/JuBEC1RW8nA?si=-a1BploFfwDJsV0a).
# Working:             ✅

# - [x] [2024.10.17] Find a way to have pre-filled values (for dates and invoice number for instance) that the user can eventually manually change or leave as it is if correct (cf.: "GeeksforGeeks - How to Set the Default Text of Tkinter Entry Widget?" (https://www.geeksforgeeks.org/how-to-set-the-default-text-of-tkinter-entry-widget/))
# - [x] [2024.10.17] Add check at the beginning of `create_invoice` to make sure that none of the fields are left empty.
# TODO: For the product to select, add a numeric field for being able to choose the quantity! This should not be part of the dictionary maybe...
# TODO: Fix the font of the invoice number, and date (maybe in the DOCX document directly)
# TODO: Complexify the script for being able to select more than 1 product -> Maybe have several DOCX templates with from 1 to 5 products -> Add a numerical field for selecting the number of products to add to the invoice, i.e., the number of product dropdown list to then DYNAMICALLY add to the UI (if possible with Tkinter!)
# TODO: Add keyboard interrupt: So that when we hit Ctrl+C from the terminal, it quits the program UI and exits. Cmd + W might work and be sufficient!
# TODO: (Finally) Improve the script to automatically populate the corresponding Excel file listing sponsor data
# TODO: "Donation" make sure to have a template extra for donations where we don't necessarily have the field "enterprise name" + we don't include TVA (-> remove the TVA part)!
# TODO: Update and generate sponsor data Excel file with data of new sponsor entered in the app here
# TODO: Make sure to have a special template for donation contracts since they should not have a secction "TVA"!

import datetime as dt
import os
import subprocess
import tkinter as tk
from datetime import datetime, timedelta
from tkinter import filedialog, messagebox

import docx
from docx2pdf import convert


class InvoiceAutomation:

    def __init__(self):
        self.root = tk.Tk()
        self.root.title("GDNC Invoice Automation")
        self.root.geometry("500x600")

        ## Labels
        # Header
        self.sponsor_company_label = tk.Label(
            self.root, text="Sponsor Company")
        self.sponsor_contact_title_label = tk.Label(
            self.root, text="Sponsor Contact Title")
        self.sponsor_contact_name_label = tk.Label(
            self.root, text="Sponsor Contact Name")
        self.sponsor_address_label = tk.Label(
            self.root, text="Sponsor Address")
        self.sponsor_postcode_label = tk.Label(
            self.root, text="Sponsor Postcode")
        self.sponsor_city_label = tk.Label(self.root, text="Sponsor City")
        # Invoice number and dates
        self.invoice_number_label = tk.Label(self.root, text="Invoice Number")
        self.issue_date_label = tk.Label(self.root, text="Issue Date")
        self.deadline_date_label = tk.Label(self.root, text="Deadline Date")
        # Products and prices
        # TODO: Here make a dictionary for the different predefined products to select from ! + maybe include the possibility to set a CUSTOM product that are not in the dropdown list and for which the user can enter its own description!
        self.product_label = tk.Label(self.root, text="Product")
        self.product_dict = {
            "Product A": {
                "description": "bla bla bla",
                "price": 123,
                "quantity": 1,
            },
            "Product B": {
                "description": "bla bla bla",
                "price": 456,
                "quantity": 2,
            },
        }

        ## Entries
        # Header
        self.sponsor_company_entry = tk.Entry(self.root)
        self.sponsor_contact_title_entry = tk.Entry(self.root)
        self.sponsor_contact_name_entry = tk.Entry(self.root)
        self.sponsor_address_entry = tk.Entry(self.root)
        self.sponsor_postcode_entry = tk.Entry(self.root)
        self.sponsor_city_entry = tk.Entry(self.root)
        # Invoice number and dates
        self.invoice_number_entry = tk.Entry(self.root, text="Invoice Number")
        self.issue_date_entry = tk.Entry(self.root, text="Issue Date")
        self.deadline_date_entry = tk.Entry(self.root, text="Deadline Date")
        # Products and prices
        # ---
        # Create a Frame to hold the Spinbox and product dropdown
        self.selection_frame = tk.Frame(self.root)
        # Create a Spinbox for quantity selection
        self.quantity_spinbox = tk.Spinbox(self.selection_frame, from_=1, to=100)
        # TODO later
        self.product = tk.StringVar(self.root)
        self.product.set("Product Number 1")
        self.product_dropdown = tk.OptionMenu(
            self.root, self.product, *self.product_dict.keys())
        # ---

        self.create_button = tk.Button(
            self.root, text="Create Invoice", command=self.create_invoice)

        padding_options = {"fill": "x", "expand": True, "padx": 5, "pady": 2}

        # Insert default values
        self.invoice_number_entry.insert(0, get_latest_invoice_number()+1)
        today = datetime.now()
        tomorrow = today + timedelta(days=1)
        tomorrow_formatted = tomorrow.strftime("%d.%m.%Y")
        in_30_days = tomorrow + timedelta(days=30)
        in_30_days_formatted = in_30_days.strftime("%d.%m.%Y")
        self.issue_date_entry.insert(0, tomorrow_formatted)
        self.deadline_date_entry.insert(0, in_30_days_formatted)

        ## Pack and add elements to the UI
        # Header
        self.sponsor_company_label.pack(padding_options)
        self.sponsor_company_entry.pack(padding_options)
        self.sponsor_contact_title_label.pack(padding_options)
        self.sponsor_contact_title_entry.pack(padding_options)
        self.sponsor_contact_name_label.pack(padding_options)
        self.sponsor_contact_name_entry.pack(padding_options)
        self.sponsor_address_label.pack(padding_options)
        self.sponsor_address_entry.pack(padding_options)
        self.sponsor_postcode_label.pack(padding_options)
        self.sponsor_postcode_entry.pack(padding_options)
        self.sponsor_city_label.pack(padding_options)
        self.sponsor_city_entry.pack(padding_options)
        # Invoice number and dates
        self.invoice_number_label.pack(padding_options)
        self.invoice_number_entry.pack(padding_options)
        self.issue_date_label.pack(padding_options)
        self.issue_date_entry.pack(padding_options)
        self.deadline_date_label.pack(padding_options)
        self.deadline_date_entry.pack(padding_options)
        # Products and prices
        # TODO: later
        self.product_label.pack(padding_options)
        # ---
        # Pack the Spinbox inside the frame
        self.quantity_spinbox.pack(side=tk.RIGHT, padx=5)
        # Pack the product dropdown inside the frame
        self.product_dropdown.pack(side=tk.RIGHT, padx=1)
        # Pack the frame
        self.selection_frame.pack(pady=10)
        # self.product_label.pack(padding_options)
        # self.product_dropdown.pack(padding_options)
        # ---

        self.entries = [
            # Header
            self.sponsor_company_entry,
            self.sponsor_contact_title_entry,
            self.sponsor_contact_name_entry,
            self.sponsor_address_entry,
            self.sponsor_postcode_entry,
            self.sponsor_city_entry,
            # Invoice number and date
            self.invoice_number_entry,
            self.issue_date_entry,
            self.deadline_date_entry,
        ]

        self.create_button.pack(padding_options)

        self.root.mainloop()

    @staticmethod
    def replace_text(paragraph, old_text, new_text):
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)

    def create_invoice(self):
        """Create invoice in PDF format.
        TODO: Update docstring!
        """
        # Check that none of the fields are left empty
        if any(entry.get().strip() for entry in self.entries):
            messagebox.showerror(title="Error", message="All fields must be completed to generate invoice!")
            return

        doc = docx.Document("invoice_models/InvoiceModel_CH95_1Entry.docx")
        
        selected_product = self.product_dict[self.product.get()]
        
        try:
            replacements = {
                # Header
                "[SPONSOR COMPANY]": self.sponsor_company_entry.get(),
                "[SPONSOR CONTACT TITLE]": self.sponsor_contact_title_entry.get(),
                "[SPONSOR CONTACT NAME]": self.sponsor_contact_name_entry.get(),
                "[SPONSOR ADDRESS]": self.sponsor_address_entry.get(),
                "[SPONSOR POSTCODE]": str(self.sponsor_postcode_entry.get()),
                "[SPONSOR CITY]": self.sponsor_city_entry.get(),
                # Invoice number and dates
                "[INVOICE NUMBER]": str(self.invoice_number_entry.get()),
                # TODO: Input automatically the date of tomorrow here!
                "[ISSUE DATE]": self.issue_date_entry.get(),
                # TODO: Input automatically the date of 1 month later minus one day here!
                "[DEADLINE DATE]": self.deadline_date_entry.get(),
                # Products and prices
                "[PRODUCT DESCRIPTION]": selected_product["description"],
                "[QUANTITY]": str(selected_product["quantity"]),
                "[PRICE]": str(selected_product["price"]),
                "[TOTAL]": f'{float(selected_product["quantity"]*selected_product["price"]):.2f}',
            }
        except ValueError:
            messagebox.showerror(title="Error", message="Invalid price!")
            return

        # Make replacements in the paragraphs of the DOCX document
        for paragraph in list(doc.paragraphs):
            for old_text, new_text in replacements.items():
                self.replace_text(paragraph=paragraph,
                                  old_text=old_text, new_text=new_text)

        # Make replacements in the tables of the DOCX document
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            self.replace_text(
                                paragraph=paragraph, old_text=old_text, new_text=new_text)

        output_docx_path = f"invoice_populated/Facture N° {str(self.invoice_number_entry.get())}"
        doc.save(output_docx_path)
        # TODO: Maybe don't use this dialog and define save_path so that we automatically save the file inside the folder "invoice_populated"
        save_path = filedialog.asksaveasfilename(
            defaultextension=".pdf", filetypes=[("PDF documents", "*.pdf")])
        # note: For converting DOCX to PDF, Microsoft Word has to be installed and opened! In addition, access to the folder in which the generated invoice will be saved has first to be granted to Microsoft Word!
        convert(input_path=output_docx_path, output_path=save_path)

        messagebox.showinfo(
            title="Success", message="Invoice created and saved successfully!")


def get_latest_invoice_number() -> str:
    """Retrieve latest invoice number from Excel file
    TODO: Update this function doctstring + real functionality
    """
    return 20240023


if __name__ == "__main__":
    InvoiceAutomation()
